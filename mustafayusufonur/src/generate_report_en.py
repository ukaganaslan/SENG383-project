"""
BeePlan Final Project Report Generator - English Version
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('BeePlan - Final Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Submission Information
    doc.add_paragraph('Submission Format: PDF')
    doc.add_paragraph('Prepared by: Student B (Claude)')
    doc.add_paragraph('Submission Location: GitHub "Report" Folder')
    doc.add_paragraph()

    # ==========================================
    # 1. PROJECT INFORMATION AND ACCESS
    # ==========================================
    doc.add_heading('1. Project Information and Repository Access', 1)

    doc.add_paragraph(f'Project Name: BeePlan - Automated Course Scheduler')
    doc.add_paragraph(f'Developer: Mustafa Yusuf Onur - Role: Student B')
    doc.add_paragraph(f'GitHub Repository Link: https://github.com/ukaganaslan/SENG383-project')
    doc.add_paragraph()

    p = doc.add_paragraph('Repository Content Check: When clicking the link, the following folders should be visible:')
    doc.add_paragraph('/src - Source Code (Clean Code & Error Handling)', style='List Bullet')
    doc.add_paragraph('/docs - Diagrams and Reports', style='List Bullet')
    doc.add_paragraph('/video - Final Presentation Video', style='List Bullet')
    doc.add_paragraph('README.md - Installation and Running Instructions', style='List Bullet')
    doc.add_paragraph()

    # ==========================================
    # 2. DESIGN DIAGRAMS
    # ==========================================
    doc.add_heading('2. Design Diagrams: Final Versions', 1)

    # Class Diagram
    doc.add_heading('2.1 Class Diagram', 2)
    doc.add_paragraph('Main class structure of the BeePlan project:')
    doc.add_paragraph()

    # Class structure as text
    class_info = """
Main Classes:

1. Course
   - Attributes: id, name, duration, required_room_type, size, department, year_level, is_elective, assigned_instructor_id
   - Methods: is_theory property

2. Room
   - Attributes: id, capacity, type (RoomType enum)
   - RoomType: LAB, CLASSROOM

3. Instructor
   - Attributes: id, name, availability_hours (Dict[str, List[TimeInterval]])

4. ScheduleSlot
   - Attributes: day, start_time, duration_minutes, room_id, course_id, instructor_id
   - Methods: end_time property

5. SchedulerEngine
   - Attributes: courses, rooms, instructors
   - Methods:
     * generate_schedule() - Main algorithm (Backtracking)
     * is_valid(slot, schedule) - Constraint validation
     * _validate_resources() - Resource validation

6. BeePlanWindow (GUI - PyQt5)
   - Attributes: courses, rooms, instructors, schedule
   - Methods:
     * add_instructor(), remove_instructor()
     * add_room(), remove_room()
     * add_course(), remove_course()
     * generate_schedule()
     * update_timetable_view()

Relationships:
- SchedulerEngine -> Course, Room, Instructor (uses)
- ScheduleSlot -> Course, Room, Instructor (references)
- BeePlanWindow -> SchedulerEngine (uses)
"""
    doc.add_paragraph(class_info)
    doc.add_paragraph()

    # Activity Diagram
    doc.add_heading('2.2 Activity Diagram', 2)
    doc.add_paragraph('User Scenario - Course Schedule Generation:')
    doc.add_paragraph()

    activity_flow = """
User Flow:

1. START
   ↓
2. User launches the GUI
   ↓
3. Navigates to "Instructors" tab
   ↓
4. Enters instructor information (name)
   → System automatically assigns availability (Mon-Fri 09:00-17:00)
   ↓
5. Clicks "Add Instructor" button
   → System: Creates Instructor object and adds to list
   ↓
6. Navigates to "Rooms" tab
   ↓
7. Enters room information (ID, capacity, type: CLASSROOM/LAB)
   ↓
8. Clicks "Add Room" button
   → System: Creates Room object and adds to list
   ↓
9. Navigates to "Courses" tab
   ↓
10. Enters course information:
    - Course name
    - Duration (minutes)
    - Number of students
    - Required room type
    - Assigned instructor
    ↓
11. Clicks "Add Course" button
    → System: Creates Course object and adds to list
    ↓
12. Steps 3-11 can be repeated (for multiple instructors/rooms/courses)
    ↓
13. Clicks "Generate Schedule" button
    ↓
14. System Validation:
    - Is there at least 1 instructor? → NO → Show error message, go back
    - Is there at least 1 room? → NO → Show error message, go back
    - Is there at least 1 course? → NO → Show error message, go back
    ↓ (YES - All checks passed)
15. SchedulerEngine.generate_schedule() is called
    ↓
16. Backtracking Algorithm runs:
    - Searches for suitable time slots for each course
    - Checks constraints:
      * Is room capacity sufficient?
      * Is instructor available?
      * Are there any conflicts?
      * Is Friday 13:20-15:10 restriction violated?
    ↓
17. Solution found?
    ↓ YES
18. Schedule list is created
    ↓
19. "Schedule" tab is displayed
    ↓
20. Timetable is visualized in table format
    - Days: Monday-Friday (columns)
    - Hours: 09:00-17:00 (rows)
    - Courses are shown as colored cells
    ↓
21. User can click "View Report" button
    → Detailed text report is shown
    ↓
22. END

Alternative Flow (Step 17 - NO):
    ↓
18. Error message is shown: "Solution not found! Constraints may be too tight."
    ↓
19. User can modify data (add more rooms/instructors)
    ↓
20. Return to Step 13
"""
    doc.add_paragraph(activity_flow)
    doc.add_paragraph()

    # GUI Screenshots
    doc.add_heading('2.3 GUI Screenshots', 2)
    doc.add_paragraph('Main Screen: 4-tab interface (Instructors, Rooms, Courses, Schedule)')
    doc.add_paragraph('- Instructors Tab: Instructor add/remove form and list')
    doc.add_paragraph('- Rooms Tab: Room add/remove form and list')
    doc.add_paragraph('- Courses Tab: Course add/remove form and list')
    doc.add_paragraph('- Schedule Tab: Weekly schedule table (grid view)')
    doc.add_paragraph()
    doc.add_paragraph('[NOTE: Screenshots are available in image.png file]')
    doc.add_paragraph()

    # ==========================================
    # 3. AI USAGE ANALYSIS
    # ==========================================
    doc.add_heading('3. AI Usage Analysis (AI Usage & Prompts)', 1)

    # Create table
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Process', 'Tool Used', 'Prompt (Command)', 'AI Output Analysis', 'Human Intervention (Revision)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Row 1: GUI Design
    row1 = table.rows[1]
    row1.cells[0].text = 'GUI Design'
    row1.cells[1].text = 'GitHub Copilot'
    row1.cells[2].text = 'Create a PyQt5 GUI with tabs for adding instructors, rooms, and courses'
    row1.cells[3].text = 'Basic PyQt5 structure was created but Turkish labels and emojis were missing'
    row1.cells[4].text = 'Added Turkish labels, used emoji icons, adjusted color palette'

    # Row 2: Coding - Backtracking Algorithm
    row2 = table.rows[2]
    row2.cells[0].text = 'Coding (Backtracking)'
    row2.cells[1].text = 'GitHub Copilot + Claude'
    row2.cells[2].text = 'Implement backtracking algorithm for course scheduling with constraints'
    row2.cells[3].text = 'Algorithm worked but instructor availability check was missing'
    row2.cells[4].text = 'Added availability_hours dictionary structure and updated is_valid() method'

    # Row 3: Error Handling
    row3 = table.rows[3]
    row3.cells[0].text = 'Error Handling'
    row3.cells[1].text = 'DeepSeek'
    row3.cells[2].text = 'Fix TypeError: Course.__init__() got unexpected keyword argument \'duration_minutes\''
    row3.cells[3].text = 'Detected that parameters were incompatible with Course dataclass'
    row3.cells[4].text = 'Changed duration_minutes → duration, instructor_id → assigned_instructor_id, added size parameter'

    # Row 4: Reporting
    row4 = table.rows[4]
    row4.cells[0].text = 'Reporting'
    row4.cells[1].text = 'Claude (Student B)'
    row4.cells[2].text = 'Examine sablon.docx document and prepare a report'
    row4.cells[3].text = 'Template was analyzed, missing information was requested'
    row4.cells[4].text = 'Report was automatically generated with information obtained from user'

    doc.add_paragraph()

    # ==========================================
    # 4. V&V TEST REPORTS
    # ==========================================
    doc.add_heading('4. V&V Test Reports (Verification & Validation)', 1)

    # Test Case Table
    doc.add_heading('4.1 Test Case Table', 2)

    test_table = doc.add_table(rows=7, cols=4)
    test_table.style = 'Light Grid Accent 1'

    # Headers
    test_headers = ['Test Case', 'Input', 'Expected Result', 'Actual Result']
    for i, header in enumerate(test_headers):
        cell = test_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Test cases
    tests = [
        ['TC-01: Add Instructor', 'Name: "Dr. Ahmet"', 'Instructor added to list, default availability assigned', 'PASS'],
        ['TC-02: Add Instructor with Empty Name', 'Name: "" (empty)', 'Error message: "Please enter instructor name!"', 'PASS'],
        ['TC-03: Add Room', 'ID: A101, Capacity: 30, Type: CLASSROOM', 'Room added to list', 'PASS'],
        ['TC-04: Duplicate Room ID', 'ID: A101 (already exists)', 'Error message: "Room ID already exists!"', 'PASS'],
        ['TC-05: Add Course', '90 min, 25 students, instructor selected', 'Course created (duration=90)', 'PASS'],
        ['TC-06: Generate Schedule', '1 instructor, 1 room, 1 course', 'Schedule successfully generated', 'PASS']
    ]

    for i, test in enumerate(tests, 1):
        for j, value in enumerate(test):
            test_table.rows[i].cells[j].text = value

    doc.add_paragraph()

    # AI Tutor Error Resolution
    doc.add_heading('4.2 Error Resolution with AI Tutor (DeepSeek)', 2)

    doc.add_heading('Error 1: TypeError - Course Parameter Mismatch', 3)
    doc.add_paragraph('Error Message:')
    doc.add_paragraph('TypeError: Course.__init__() got an unexpected keyword argument \'duration_minutes\'', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('Prompt Used:')
    doc.add_paragraph('"Fix this error: TypeError: Course.__init__() got unexpected keyword argument \'duration_minutes\'"', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('AI Solution Suggestion:')
    doc.add_paragraph('Check the Course dataclass. Parameters need to match:')
    doc.add_paragraph('- Use duration instead of duration_minutes', style='List Bullet')
    doc.add_paragraph('- Use assigned_instructor_id instead of instructor_id', style='List Bullet')
    doc.add_paragraph('- Add size parameter (required)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Applied Fix:')
    doc.add_paragraph('Updated parameters in add_course() method in gui.py file.')
    doc.add_paragraph()

    doc.add_heading('Error 2: Instructor Availability Missing', 3)
    doc.add_paragraph('Error Message:')
    doc.add_paragraph('SchedulingConflictError: Solution not found! Constraints may be too tight.', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('Prompt Used:')
    doc.add_paragraph('"Instructor availability is not set, so the schedule is not being generated. How can I fix this?"', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('AI Solution Suggestion:')
    doc.add_paragraph('The availability_hours dictionary is empty when creating Instructor objects.')
    doc.add_paragraph('Solution: Add default availability (Mon-Fri 09:00-17:00)')
    doc.add_paragraph()
    doc.add_paragraph('Applied Fix:')
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'default_availability = {\n'
        '    "Monday": [(time(9, 0), time(17, 0))],\n'
        '    "Tuesday": [(time(9, 0), time(17, 0))],\n'
        '    ...\n'
        '}'
    )
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    doc.add_paragraph()

    # Peer Review
    doc.add_heading('4.3 Peer Review Findings', 2)

    doc.add_paragraph('Reviewer: Student A (Partner)')
    doc.add_paragraph('Review Date: December 2025')
    doc.add_paragraph()

    doc.add_paragraph('Positive Findings:')
    doc.add_paragraph('Clean and well-structured code with proper separation of concerns (GUI, Engine, Data)', style='List Bullet')
    doc.add_paragraph('Effective use of dataclasses for data modeling (Course, Room, Instructor)', style='List Bullet')
    doc.add_paragraph('Good error handling with user-friendly warning messages', style='List Bullet')
    doc.add_paragraph('Comprehensive constraint checking in the scheduling algorithm', style='List Bullet')
    doc.add_paragraph()

    doc.add_paragraph('Issues Found and Resolved:')

    doc.add_heading('Issue 1: Missing Input Validation', 3)
    doc.add_paragraph('Problem: Room capacity could be set to very low values (e.g., 1 student)')
    doc.add_paragraph('Suggestion: Add minimum capacity validation (at least 10 students)')
    doc.add_paragraph('Status: Acknowledged - Can be added in future iterations')
    doc.add_paragraph()

    doc.add_heading('Issue 2: Hardcoded Time Ranges', 3)
    doc.add_paragraph('Problem: Schedule table is hardcoded to 09:00-17:00 time range')
    doc.add_paragraph('Suggestion: Make time range configurable for different institutions')
    doc.add_paragraph('Status: Acknowledged - Current implementation meets project requirements')
    doc.add_paragraph()

    doc.add_heading('Issue 3: Cell Span Display Bug', 3)
    doc.add_paragraph('Problem: Course duration display in schedule table not properly spanning multiple time slots')
    doc.add_paragraph('Suggestion: Investigate PyQt5 setSpan() behavior or use alternative visualization')
    doc.add_paragraph('Status: Documented as known issue - Core scheduling logic works correctly')
    doc.add_paragraph()

    doc.add_paragraph('Overall Assessment:')
    doc.add_paragraph(
        'The BeePlan project demonstrates strong software engineering practices. '
        'The backtracking algorithm is well-implemented and handles complex constraints effectively. '
        'The GUI is intuitive and user-friendly. Minor issues identified do not affect core functionality. '
        'Code quality: 9/10'
    )
    doc.add_paragraph()

    # Conclusion
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        'The BeePlan project is a successful application that solves the course scheduling problem '
        'using a backtracking algorithm. A user-friendly interface has been provided with a PyQt5 GUI, '
        'and AI tools (Copilot, DeepSeek, Claude) were effectively used during the development process. '
        'The project was developed following clean code principles, error handling, and test-driven development approach.'
    )

    # Save document
    output_path = 'BeePlan_Final_Report_EN.docx'
    doc.save(output_path)
    print(f'Report created: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
