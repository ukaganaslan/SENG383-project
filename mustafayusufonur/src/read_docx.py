from docx import Document
import sys

# UTF-8 encoding için
sys.stdout.reconfigure(encoding='utf-8')

doc = Document('sablon.docx')

print('=== SABLON.DOCX İÇERİĞİ ===\n')
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'{i+1}. {para.text}')

# Tabloları da kontrol et
if doc.tables:
    print('\n=== TABLOLAR ===\n')
    for table_idx, table in enumerate(doc.tables):
        print(f'\nTablo {table_idx + 1}:')
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            print(f'  Satır {row_idx + 1}: {" | ".join(row_data)}')
