"""
BeePlan Final Project Report Generator
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Başlık
    title = doc.add_heading('BeePlan - Final Proje Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Teslim Bilgileri
    doc.add_paragraph('Teslim Formatı: PDF')
    doc.add_paragraph('Hazırlama Aracı: Student B (Claude)')
    doc.add_paragraph('Teslim Yeri: GitHub "Report" Klasörü')
    doc.add_paragraph()

    # ==========================================
    # 1. PROJE KÜNYESI VE ERİŞİM
    # ==========================================
    doc.add_heading('1. Proje Künyesi ve Erişim (Repository Access)', 1)

    doc.add_paragraph(f'Proje Adı: BeePlan - Automated Course Scheduler')
    doc.add_paragraph(f'Geliştirici: Mustafa Yusuf Onur - Rol: Student B')
    doc.add_paragraph(f'GitHub Repository Linki: https://github.com/ukaganaslan/SENG383-project')
    doc.add_paragraph()

    p = doc.add_paragraph('Repo İçeriği Kontrolü: Linke tıklandığında aşağıdaki klasörler görünmelidir:')
    doc.add_paragraph('/src - Kaynak Kodlar (Clean Code & Error Handling)', style='List Bullet')
    doc.add_paragraph('/docs - Diyagramlar ve Raporlar', style='List Bullet')
    doc.add_paragraph('/video - Final Presentation Video', style='List Bullet')
    doc.add_paragraph('README.md - Kurulum ve Çalıştırma Talimatları', style='List Bullet')
    doc.add_paragraph()

    # ==========================================
    # 2. TASARIM DİYAGRAMLARI
    # ==========================================
    doc.add_heading('2. Tasarım Diyagramları: Son Versiyonlar', 1)

    # Class Diagram
    doc.add_heading('2.1 Class Diagram (Sınıf Diyagramı)', 2)
    doc.add_paragraph('BeePlan projesinin ana sınıf yapısı:')
    doc.add_paragraph()

    # Class yapısını metin olarak ekle
    class_info = """
Ana Sınıflar:

1. Course (Ders)
   - Attributes: id, name, duration, required_room_type, size, department, year_level, is_elective, assigned_instructor_id
   - Methods: is_theory property

2. Room (Sınıf/Oda)
   - Attributes: id, capacity, type (RoomType enum)
   - RoomType: LAB, CLASSROOM

3. Instructor (Hoca)
   - Attributes: id, name, availability_hours (Dict[str, List[TimeInterval]])

4. ScheduleSlot (Program Zaman Dilimi)
   - Attributes: day, start_time, duration_minutes, room_id, course_id, instructor_id
   - Methods: end_time property

5. SchedulerEngine (Çizelgeleme Motoru)
   - Attributes: courses, rooms, instructors
   - Methods:
     * generate_schedule() - Ana algoritma (Backtracking)
     * is_valid(slot, schedule) - Kısıtlama kontrolü
     * _validate_resources() - Kaynak doğrulama

6. BeePlanWindow (GUI - PyQt5)
   - Attributes: courses, rooms, instructors, schedule
   - Methods:
     * add_instructor(), remove_instructor()
     * add_room(), remove_room()
     * add_course(), remove_course()
     * generate_schedule()
     * update_timetable_view()

İlişkiler:
- SchedulerEngine -> Course, Room, Instructor (uses)
- ScheduleSlot -> Course, Room, Instructor (references)
- BeePlanWindow -> SchedulerEngine (uses)
"""
    doc.add_paragraph(class_info)
    doc.add_paragraph()

    # Activity Diagram
    doc.add_heading('2.2 Activity Diagram (Aktivite Diyagramı)', 2)
    doc.add_paragraph('Kullanıcı Senaryosu - Ders Programı Oluşturma:')
    doc.add_paragraph()

    activity_flow = """
Kullanıcı Akışı:

1. START
   ↓
2. Kullanıcı GUI'yi başlatır
   ↓
3. "Instructors" sekmesine gider
   ↓
4. Hoca bilgilerini girer (isim)
   → Sistem otomatik müsaitlik atar (Pzt-Cuma 09:00-17:00)
   ↓
5. "Add Instructor" butonuna tıklar
   → Sistem: Instructor nesnesi oluşturur ve listeye ekler
   ↓
6. "Rooms" sekmesine gider
   ↓
7. Sınıf bilgilerini girer (ID, kapasite, tip: CLASSROOM/LAB)
   ↓
8. "Add Room" butonuna tıklar
   → Sistem: Room nesnesi oluşturur ve listeye ekler
   ↓
9. "Courses" sekmesine gider
   ↓
10. Ders bilgilerini girer:
    - Ders adı
    - Süre (dakika)
    - Öğrenci sayısı
    - Oda tipi gereksinimi
    - Atanan hoca
    ↓
11. "Add Course" butonuna tıklar
    → Sistem: Course nesnesi oluşturur ve listeye ekler
    ↓
12. Adım 3-11 tekrarlanabilir (birden fazla hoca/sınıf/ders için)
    ↓
13. "Generate Schedule" butonuna tıklar
    ↓
14. Sistem Kontrolü:
    - En az 1 instructor var mı? → HAYIR → Hata mesajı göster, geri dön
    - En az 1 room var mı? → HAYIR → Hata mesajı göster, geri dön
    - En az 1 course var mı? → HAYIR → Hata mesajı göster, geri dön
    ↓ (EVET - Tüm kontroller geçti)
15. SchedulerEngine.generate_schedule() çağrılır
    ↓
16. Backtracking Algoritması çalışır:
    - Her ders için uygun zaman dilimi arar
    - Kısıtlamaları kontrol eder:
      * Oda kapasitesi yeterli mi?
      * Hoca müsait mi?
      * Çakışma var mı?
      * Cuma 13:20-15:10 yasağı ihlal edilmiş mi?
    ↓
17. Çözüm bulundu mu?
    ↓ EVET
18. Schedule listesi oluşturulur
    ↓
19. "Schedule" sekmesi gösterilir
    ↓
20. Çizelge tabloda görselleştirilir
    - Günler: Pazartesi-Cuma (sütunlar)
    - Saatler: 09:00-17:00 (satırlar)
    - Dersler renkli hücreler olarak gösterilir
    ↓
21. Kullanıcı "View Report" butonuna tıklayabilir
    → Detaylı metin raporu gösterilir
    ↓
22. END

Alternatif Akış (Adım 17 - HAYIR):
    ↓
18. Hata mesajı gösterilir: "Çözüm bulunamadı! Kısıtlamalar çok sıkı olabilir."
    ↓
19. Kullanıcı verileri düzenleyebilir (daha fazla sınıf/hoca ekler)
    ↓
20. Adım 13'e geri dön
"""
    doc.add_paragraph(activity_flow)
    doc.add_paragraph()

    # GUI Screenshots
    doc.add_heading('2.3 GUI Screenshots', 2)
    doc.add_paragraph('Ana Ekran: 4 sekmeli arayüz (Instructors, Rooms, Courses, Schedule)')
    doc.add_paragraph('- Instructors Tab: Hoca ekleme/silme formu ve liste')
    doc.add_paragraph('- Rooms Tab: Sınıf ekleme/silme formu ve liste')
    doc.add_paragraph('- Courses Tab: Ders ekleme/silme formu ve liste')
    doc.add_paragraph('- Schedule Tab: Haftalık program tablosu (grid view)')
    doc.add_paragraph()
    doc.add_paragraph('[NOT: Screenshots image.png dosyasında mevcuttur]')
    doc.add_paragraph()

    # ==========================================
    # 3. AI KULLANIM ANALİZİ
    # ==========================================
    doc.add_heading('3. AI Kullanım Analizi (AI Usage & Prompts)', 1)

    # Tablo oluştur
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Light Grid Accent 1'

    # Başlıklar
    headers = ['Süreç', 'Kullanılan Araç', 'Prompt (Komut)', 'AI Çıktı Analizi', 'İnsan Müdahalesi (Revision)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Satır 1: GUI Tasarımı
    row1 = table.rows[1]
    row1.cells[0].text = 'GUI Tasarımı'
    row1.cells[1].text = 'GitHub Copilot'
    row1.cells[2].text = 'Create a PyQt5 GUI with tabs for adding instructors, rooms, and courses'
    row1.cells[3].text = 'Temel PyQt5 yapısı oluşturuldu ancak Türkçe label\'lar ve emoji\'ler eksikti'
    row1.cells[4].text = 'Türkçe etiketler eklendi, emoji ikonları kullanıldı, renk paleti düzenlendi'

    # Satır 2: Kodlama - Backtracking Algoritması
    row2 = table.rows[2]
    row2.cells[0].text = 'Kodlama (Backtracking)'
    row2.cells[1].text = 'GitHub Copilot + Claude'
    row2.cells[2].text = 'Implement backtracking algorithm for course scheduling with constraints'
    row2.cells[3].text = 'Algoritma çalıştı ama hoca müsaitlik kontrolü eksikti'
    row2.cells[4].text = 'availability_hours dictionary yapısı eklendi ve is_valid() metodu güncellendi'

    # Satır 3: Error Handling
    row3 = table.rows[3]
    row3.cells[0].text = 'Error Handling'
    row3.cells[1].text = 'DeepSeek'
    row3.cells[2].text = 'Fix TypeError: Course.__init__() got unexpected keyword argument \'duration_minutes\''
    row3.cells[3].text = 'Parametrelerin Course dataclass ile uyumsuz olduğunu tespit etti'
    row3.cells[4].text = 'duration_minutes → duration, instructor_id → assigned_instructor_id, size parametresi eklendi'

    # Satır 4: Raporlama
    row4 = table.rows[4]
    row4.cells[0].text = 'Raporlama'
    row4.cells[1].text = 'Claude (Student B)'
    row4.cells[2].text = 'sablon.docx dökünmanını incele bir rapor hazırlayacağız'
    row4.cells[3].text = 'Şablon analiz edildi, eksik bilgiler soruldu'
    row4.cells[4].text = 'Kullanıcıdan alınan bilgilerle rapor otomatik oluşturuldu'

    doc.add_paragraph()

    # ==========================================
    # 4. V&V TEST RAPORLARI
    # ==========================================
    doc.add_heading('4. V&V Test Raporları (Verification & Validation)', 1)

    # Test Case Tablosu
    doc.add_heading('4.1 Test Case Tablosu', 2)

    test_table = doc.add_table(rows=7, cols=4)
    test_table.style = 'Light Grid Accent 1'

    # Başlıklar
    test_headers = ['Test Case', 'Input', 'Beklenen Sonuç', 'Gerçek Sonuç']
    for i, header in enumerate(test_headers):
        cell = test_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Test case'ler
    tests = [
        ['TC-01: Hoca Ekleme', 'İsim: "Dr. Ahmet"', 'Hoca listeye eklenir, varsayılan müsaitlik atanır', '✓ Başarılı'],
        ['TC-02: Boş İsimle Hoca Ekleme', 'İsim: "" (boş)', 'Hata mesajı: "Please enter instructor name!"', '✓ Başarılı'],
        ['TC-03: Sınıf Ekleme', 'ID: A101, Kapasite: 30, Tip: CLASSROOM', 'Sınıf listeye eklenir', '✓ Başarılı'],
        ['TC-04: Duplicate Sınıf ID', 'ID: A101 (zaten var)', 'Hata mesajı: "Room ID already exists!"', '✓ Başarılı'],
        ['TC-05: Ders Ekleme', '90 dk, 25 öğrenci, hoca seçili', 'Ders oluşturulur (duration=90)', '✓ Başarılı'],
        ['TC-06: Program Oluşturma', '1 hoca, 1 sınıf, 1 ders', 'Çizelge başarıyla oluşturulur', '✓ Başarılı']
    ]

    for i, test in enumerate(tests, 1):
        for j, value in enumerate(test):
            test_table.rows[i].cells[j].text = value

    doc.add_paragraph()

    # AI Tutor ile Hata Çözümü
    doc.add_heading('4.2 AI Tutor ile Hata Çözümü (DeepSeek)', 2)

    doc.add_heading('Hata 1: TypeError - Course Parameter Mismatch', 3)
    doc.add_paragraph('Hata Mesajı:')
    doc.add_paragraph('TypeError: Course.__init__() got an unexpected keyword argument \'duration_minutes\'', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('Kullanılan Prompt:')
    doc.add_paragraph('"Fix this error: TypeError: Course.__init__() got unexpected keyword argument \'duration_minutes\'"', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('AI Çözüm Önerisi:')
    doc.add_paragraph('Course dataclass\'ını kontrol et. Parametrelerin eşleşmesi gerekiyor:')
    doc.add_paragraph('- duration_minutes yerine duration kullan', style='List Bullet')
    doc.add_paragraph('- instructor_id yerine assigned_instructor_id kullan', style='List Bullet')
    doc.add_paragraph('- size parametresini ekle (zorunlu)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Uygulanan Düzeltme:')
    doc.add_paragraph('gui.py dosyasındaki add_course() metodunda parametreler güncellendi.')
    doc.add_paragraph()

    doc.add_heading('Hata 2: Instructor Availability Missing', 3)
    doc.add_paragraph('Hata Mesajı:')
    doc.add_paragraph('SchedulingConflictError: Çözüm bulunamadı! Kısıtlamalar çok sıkı olabilir.', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('Kullanılan Prompt:')
    doc.add_paragraph('"Hocaların müsaitliği ayarlanmadığı için program oluşturmuyor, nasıl düzeltebilirim?"', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('AI Çözüm Önerisi:')
    doc.add_paragraph('Instructor nesnesi oluşturulurken availability_hours dictionary\'si boş kalıyor.')
    doc.add_paragraph('Çözüm: Varsayılan müsaitlik ekle (Pzt-Cuma 09:00-17:00)')
    doc.add_paragraph()
    doc.add_paragraph('Uygulanan Düzeltme:')
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'default_availability = {\n'
        '    "Monday": [(time(9, 0), time(17, 0))],\n'
        '    "Tuesday": [(time(9, 0), time(17, 0))],\n'
        '    ...\n'
        '}'
    )
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    doc.add_paragraph()

    # Peer Review
    doc.add_heading('4.3 Peer Review Bulguları', 2)

    doc.add_paragraph('İncelemeci: Student A (Partner)')
    doc.add_paragraph('İnceleme Tarihi: Aralık 2025')
    doc.add_paragraph()

    doc.add_paragraph('Olumlu Bulgular:')
    doc.add_paragraph('Temiz ve iyi yapılandırılmış kod (GUI, Engine, Data ayrımı)', style='List Bullet')
    doc.add_paragraph('Veri modellemesi için dataclass kullanımı etkili (Course, Room, Instructor)', style='List Bullet')
    doc.add_paragraph('Kullanıcı dostu hata mesajları ile iyi hata yönetimi', style='List Bullet')
    doc.add_paragraph('Çizelgeleme algoritmasında kapsamlı kısıt kontrolü', style='List Bullet')
    doc.add_paragraph()

    doc.add_paragraph('Bulunan ve Çözülen Sorunlar:')

    doc.add_heading('Sorun 1: Eksik Girdi Doğrulama', 3)
    doc.add_paragraph('Problem: Sınıf kapasitesi çok düşük değerlere ayarlanabiliyor (örn: 1 öğrenci)')
    doc.add_paragraph('Öneri: Minimum kapasite doğrulaması ekle (en az 10 öğrenci)')
    doc.add_paragraph('Durum: Kabul edildi - Gelecek sürümlerde eklenebilir')
    doc.add_paragraph()

    doc.add_heading('Sorun 2: Sabit Kodlanmış Zaman Aralıkları', 3)
    doc.add_paragraph('Problem: Program tablosu 09:00-17:00 zaman aralığına sabit kodlanmış')
    doc.add_paragraph('Öneri: Farklı kurumlar için zaman aralığını yapılandırılabilir yap')
    doc.add_paragraph('Durum: Kabul edildi - Mevcut uygulama proje gereksinimlerini karşılıyor')
    doc.add_paragraph()

    doc.add_heading('Sorun 3: Hücre Birleştirme Görüntüleme Hatası', 3)
    doc.add_paragraph('Problem: Program tablosunda ders süresi gösterimi birden fazla zaman dilimini düzgün kaplamıyor')
    doc.add_paragraph('Öneri: PyQt5 setSpan() davranışını araştır veya alternatif görselleştirme kullan')
    doc.add_paragraph('Durum: Bilinen sorun olarak belgelendi - Temel çizelgeleme mantığı doğru çalışıyor')
    doc.add_paragraph()

    doc.add_paragraph('Genel Değerlendirme:')
    doc.add_paragraph(
        'BeePlan projesi güçlü yazılım mühendisliği uygulamalarını göstermektedir. '
        'Backtracking algoritması iyi uygulanmış ve karmaşık kısıtlamaları etkili şekilde ele almaktadır. '
        'GUI sezgisel ve kullanıcı dostudur. Belirlenen küçük sorunlar temel işlevselliği etkilememektedir. '
        'Kod kalitesi: 9/10'
    )
    doc.add_paragraph()

    # Sonuç
    doc.add_heading('Sonuç', 1)
    doc.add_paragraph(
        'BeePlan projesi, ders çizelgeleme problemini backtracking algoritması kullanarak '
        'çözen başarılı bir uygulamadır. PyQt5 GUI ile kullanıcı dostu bir arayüz sağlanmış, '
        'AI araçları (Copilot, DeepSeek, Claude) geliştirme sürecinde etkin kullanılmıştır. '
        'Proje, clean code prensipleri, error handling, ve test-driven development yaklaşımı '
        'ile geliştirilmiştir.'
    )

    # Dosyayı kaydet
    output_path = 'BeePlan_Final_Report.docx'
    doc.save(output_path)
    print(f'Rapor olusturuldu: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
